import os,re,torch
import pypandoc
from dotenv import load_dotenv
load_dotenv()
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api.formatters import TextFormatter
from sentence_transformers import SentenceTransformer
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser
from langchain_groq import ChatGroq
from docx import Document
from youtube_transcript_api import YouTubeTranscriptApi
from youtube_transcript_api.proxies import WebshareProxyConfig
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import subprocess
from docx.enum.style import WD_STYLE_TYPE
import sounddevice as sd
import numpy as npe
from scipy.io.wavfile import write


ytt_api = YouTubeTranscriptApi(
    proxy_config=WebshareProxyConfig(
        proxy_username="<bgukzzwo>",
        proxy_password="<4kbw1ocmr69l>",
    )
)

def TranscriptExtracter(Url):
	CorrectedUrl=re.search(r'(?:youtu\.be/|v=|embed/)([A-Za-z0-9_-]{11})', Url).group(1)
	VideoId=CorrectedUrl
	Transcript=YouTubeTranscriptApi().fetch(VideoId, languages=['en'], preserve_formatting=True)
	Text = TextFormatter().format_transcript(Transcript)
	return Text
# ----------------------------------------
# 2. Build Vectorstore + Retriever
# ----------------------------------------
def build_vectorstore(text):
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")

    chunks = [text[i:i + 1200] for i in range(0, len(text), 1200)]

    vectorstore = FAISS.from_texts(chunks, embedding=embeddings)
    return vectorstore.as_retriever()

MASTER_PROMPT="""
You are a senior non-fiction book author, editor, and writing strategist with 20+ years of experience.

Your task is to take the user-provided YouTube transcript and transform it into a complete,
well-structured, professionally written non-fiction BOOK while keeping the content of the transcript.

Follow these instructions EXACTLY:

------------------
### 1. BOOK STRUCTURE
Create a full-length book with:

- A compelling **book title** written with a heading format: Title: [Title]
- A relevant **subtitle** written with a heading Subtitle
- A short but powerful **book description** with heading Book Description
- A **preface** section with heading Preface
- A detailed **table of contents** in form of bibliographic entries without writing Chapter instead just write the chapter title with proper numbering and sub-heading numbers inside an indendation of each chapter number and it should not contain the page number
- exclude any mention of "conclusion" or "resource" in the table of contents
- 8–12 **chapters**, depending on transcript depth
- Each chapter must include:
  - A chapter title in the form "Chapter X: [Title]"
  - An engaging introduction with heading "Introduction"
  - Multiple sub-sections with headings in format "Sub‑section Chapter X.Y: [Sub-section Title]"
  - Examples, stories, analogies, or metaphors (even if not in transcript)
  - Clear explanations & insights

- A **final conclusion** with heading Conclusion
- A **resources or further reading** section (optional) wiht heading Resources or Further Reading

------------------
### 2. WRITING STYLE GUIDELINES
- Write in **clear, engaging, beginner-friendly** language.
- Maintain the speaker’s original message but rewrite in book-quality prose.
- Remove filler words, repetition, and irrelevant content.
- Expand ideas logically — it should read like a book, not a transcript.
- Add explanation where needed so readers understand even without watching the video.
- Improve flow, transitions, examples, and storytelling.
- Make the tone **professional**, **insightful**, and **engaging**.

------------------
### 3. CONTENT TRANSFORMATION RULES
- Reorganize ideas from the transcript into a logical book structure.
- Combine scattered points into coherent chapters.
- Expand short ideas into complete explanations.
- Add new insights where necessary (but stay faithful to the topic).
- Remove timestamps, speaker tags, YouTube artifacts, or formatting.
- Fix grammar, clarity, and structure.
- The should be no space between chapter ending and next chapter heading.


Do NOT format paragraphs using bullet points, hyphens, or list markers. 
Write in normal paragraph style, not as a list. 
No sentence should begin with "-" or "*" or "•" unless explicitly asked.


------------------
### 4. QUALITY REQUIREMENTS
The final result must be:

- Fully polished and ready for publishing
- Human-like, rich, and detailed
- Long enough to feel like a proper book (8–25k words depending on model)
- Coherent, structured, and logically flowing
- Free from transcript-style repetitions

------------------
### 5. CONTEXT
Here is the transcript context. Use it deeply to construct the book:

{context}


Return ONLY the final book. Do not explain your process.

    """
# ----------------------------------------
# 3. Create RAG Pipeline (New LangChain 0.3+)
# ----------------------------------------
def create_rag_chain(retriever):
    prompt = ChatPromptTemplate.from_template(MASTER_PROMPT)

    llm = ChatGroq(model="openai/gpt-oss-120b",temperature=0.6,max_retries=3,api_key=os.getenv("GROQ_API_KEY"))

    rag_chain = (
        {"context": retriever}
        | prompt
        | llm
        | StrOutputParser()
    )

    return rag_chain

# Get response according to the transcript


# ----------------------------------------
# 4. Full Book Generator
# ----------------------------------------
def generate_book(text):
    retriever = build_vectorstore(text)
    rag_chain = create_rag_chain(retriever)

    return rag_chain.invoke(MASTER_PROMPT)

# ----------------------------------------
# 5. Save to Word Document
# ----------------------------------------

def clean_output(text):
    # 1. Remove leading "- ", "* ", "• "
    text = re.sub(r"^\s*[-\*\•]\s*", "", text, flags=re.MULTILINE)

    # 2. Remove leading bold markers like "**Example:**"
    text = re.sub(r"^\s*\*\*(.*?)\*\*\s*:? ?", r"\1: ", text, flags=re.MULTILINE)

    # 3. Remove bold markers WITHOUT replacing them with text (optional)
    # text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)

    # 4. Cleanup extra spaces and blank lines
    text = re.sub(r"\n\s*\n\s*\n+", "\n\n", text)

    return text.strip()


def save_as_word(full_text, filename="generated_book.docx"):
    doc = Document()
    author="GENERATED BY: QUAD AI"
    styles = doc.styles
    if "AuthorStyle" not in styles:
          author_style = styles.add_style("AuthorStyle", WD_STYLE_TYPE.PARAGRAPH)
          author_style.base_style = styles["Normal"]
    # -----------------------------
    # CLEAN & SPLIT RAW TEXT
    # -----------------------------
    lines = full_text.split("\n")

    cleaned = []
    for line in lines:
        # Remove markdown symbols
        line = line.replace("**", "")
        line = line.replace("###", "")
        line = line.replace("##", "")
        line = line.replace("#", "")
        line = line.lstrip("-*• ")  # remove bullets

        if line.strip():
            cleaned.append(line.strip())

    # -----------------------------
    # EXTRACT TITLE & SUBTITLE
    # -----------------------------
    title = cleaned[0][6:] if len(cleaned) > 0 else "Untitled Book"
    subtitle = cleaned[1][9:] if len(cleaned) > 1 else ""
    # Remove title + subtitle from body text
    body_lines = cleaned[2:] if len(cleaned) > 2 else []

    # -----------------------------
    # TITLE PAGE (CENTERED, STYLED)
    # -----------------------------
    # add vertical space to center the text
    for _ in range(10):
        doc.add_paragraph("")

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.style ="Title"

    title_run = title_para.add_run(title)
    title_run.bold = True
    title_run.font.size = Pt(36)
    title_run.font.name = "Georgia"

    if subtitle:
        doc.add_paragraph("")
        subtitle_para = doc.add_paragraph()
        subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_para.style ="Subtitle"
        sub_run = subtitle_para.add_run(subtitle)
        sub_run.font.size = Pt(20)
        sub_run.font.italic = True
        sub_run.font.name = "Georgia"
        
    author_para = doc.add_paragraph()
    author_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author_para.style ="AuthorStyle"
    author_run = author_para.add_run(author)
    author_run.font.size = Pt(18)		  

    doc.add_page_break()

    # -----------------------------
    # MAIN CONTENT (CHAPTERS + TEXT)
    # -----------------------------
    first_chapter = True

    for line in body_lines:

        # Detect chapter heading
        if line.startswith("Chapter"):
            doc.add_page_break()
            first_chapter = True
            heading_para = doc.add_heading(line, level=1)
            heading_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run=heading_para.runs[0]
            run.bold = True
            run.font.size = Pt(18)
            run.font.name = "Georgia"
            doc.add_paragraph("")  # spacing
        elif line.lower().startswith("table of contents"):
            doc.add_page_break()
            toc_heading = doc.add_heading("Table of Contents", level=1)
            toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run=toc_heading.runs[0]
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Georgia"
            doc.add_paragraph("")
        elif line.lower().startswith("preface") or line.lower().startswith("book description"):
            preface_heading = doc.add_heading(line, level=2)
            preface_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run=preface_heading.runs[0]
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Georgia"            
            doc.add_paragraph("")  # spacing
        elif line.lower().startswith("sub‑section") or line.lower().startswith("sub-section"):
            sub_heading = doc.add_heading(line[20:], level=3)
            sub_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run=sub_heading.runs[0]
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Georgia"            
            doc.add_paragraph("")  # spacing
        elif line.lower().startswith("introduction"):
            intro_heading = doc.add_heading(line, level=2)
            intro_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run=intro_heading.runs[0]
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Georgia"            
            doc.add_paragraph("")  # spacing
        elif line.lower().startswith("conclusion"):
            doc.add_page_break()
            conclusion_heading = doc.add_heading(line, level=1)
            conclusion_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run=conclusion_heading.runs[0]
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Georgia"            
            doc.add_paragraph("")  # spacing 
        elif line.lower().startswith("resources") or line.lower().startswith("further reading"):
            conclusion_heading = doc.add_heading(line, level=1)
            conclusion_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run=conclusion_heading.runs[0]
            run.bold = True
            run.font.size = Pt(14)
            run.font.name = "Georgia"
            doc.add_paragraph("")  # spacing
        else:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(12)
            run.font.name = "Georgia"
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # -----------------------------
    # SAVE FILE
    # -----------------------------
    doc.save(filename)
    print(f"\nWord book created: {filename}")
    return filename

def docx_to_pdf(input_file, output_folder=None):
    """
    Convert a DOCX file to PDF using LibreOffice (headless mode).
    """
    # If no output folder is provided, use the same directory
    if output_folder is None:
        output_folder = os.path.dirname(os.path.abspath(input_file))

    # LibreOffice installation path (default Windows path)
    libreoffice_path = r"C:\Program Files\LibreOffice\program\soffice.exe"

    command = [
        libreoffice_path,
        "--headless",
        "--convert-to", "pdf",
        "--outdir", output_folder,
        input_file
    ]

    print("Running:", " ".join(command))
    subprocess.run(command, check=True)

    pd_filename = os.path.splitext(os.path.basename(input_file))[0] + ".pdf"
    print(f"PDF file created: {pd_filename}")
    return pd_filename

